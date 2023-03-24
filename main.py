# pip install speechrecognition
# pip install pyaudio
# pip install pywhatkit
# pip install TTsx3
# pip install gtts
# pip install googlesearch-python


import speech_recognition as sr
import pywhatkit
from gtts import gTTS
import pygame
import os
import pyttsx3
import datetime
import googlesearch
import pyjokes
import requests
from bs4 import BeautifulSoup
from geopy.geocoders import Nominatim
import json
import webbrowser
from PIL import Image
import pyowm
import wikipedia
import geocoder
import sys
import docx
import time
import smtplib
import spotipy
from spotipy.oauth2 import spotifyOAuth
from twilio.rest import Client
import subprocess
import openai
from transformers import pipeline
import pytz
from google.oauth2 import service_account
from googleapiclient.discovery import build
import boto3
import spacy
from googletrans import Translator




#speech engine initialization
engine = pyttsx3.init()
voices = engine.getProperty("voices")
engine.setProperty("voices", voices[1].id,)
activationword = "rafiki"

def speak(text, rate = 150):
    engine.setProperty("rate",rate)
    engine.say(text)
    engine.runandwait()

def get_audio():
    recorder = sr.Recognizer()
    with sr.Microphone() as source:
        print("nasikiliza...")
        recorder.pause_threshold = 2
        audio = recorder.listen(source)

    text = recorder.recognize_google(audio)
    print(text)


    if "rafiki" in text.lower():
        play_audio(r"C:\Users\admin\Desktop\beep-07a.wav")

    print(f"you said: {text}")
    return text

# Create NLP objects for each language
nlp_kiswahili = spacy.load("xx_ent_wiki_sm") # Kiswahili
nlp_zulu = spacy.load("xx_ent_wiki_sm") # Zulu
nlp_hausa = spacy.load("xx_ent_wiki_sm") # Hausa
nlp_masri = spacy.load("xx_ent_wiki_sm") # Masri (Egyptian Arabic)
nlp_ladi = spacy.load("xx_ent_wiki_sm") # Ladi

# Create a dictionary mapping language codes to NLP objects
nlp_dict = {
    "kiswahili": nlp_kiswahili,
    "zulu": nlp_zulu,
    "hausa": nlp_hausa,
    "masri": nlp_masri,
    "ladi": nlp_ladi
}

# Create a Translator object for translating between languages
translator = Translator()

# Function to detect language of text and return NLP object
def detect_language(text):
    # Use Google Translate API to detect language
    lang = translator.detect(text).lang
    # Return NLP object for detected language
    return nlp_dict.get(lang, None)

# Function to translate text from one language to another
def translate(text, target_lang):
    # Use Google Translate API to translate text
    translation = translator.translate(text, dest=target_lang)
    return translation.text

def speech(text,):
    # initialize tts, create mp3 and save
    print(text)
    language = "sw"
    output_file = "output.mp3"
    tts = gTTS(text=text, lang=language, slow=False)
    tts.save(output_file)

    # play the audio file using pygame
    is_playing = play_audio(output_file)
    while is_playing:
        is_playing = pygame.mixer.music.get_busy()

    # stop the audio playback and release resources
    pygame.mixer.music.stop()
    pygame.mixer.quit()

    # remove the audio file
    os.remove(output_file)


def play_audio(filename):
    pygame.mixer.init()
    pygame.mixer.music.load(filename)
    pygame.mixer.music.play()
    while pygame.mixer.music.get_busy():
        continue
    return True  # Return true when the audio is finished playing

# Define a function to speak the current time

# Set up OpenAI API credentials
openai.api_key = os.getenv("OPENAI_API_KEY")

# Define the question and context
question = "What is the capital city of Japan?"
context = "Japan is an island country in East Asia, located in the Pacific Ocean. Its capital city is Tokyo."

# Use OpenAI's GPT-3 model to generate answer
response = openai.Completion.create(
    engine="davinci", # specify which GPT-3 model to use
    prompt=f"Q: {question}\nA: {context}",
    max_tokens=1024,
)

# Print the answer
answer = response.choices[0].text
print(answer)

# Define the question and context
question = "What is the capital city of Japan?"
context = "Japan is an island country in East Asia, located in the Pacific Ocean. Its capital city is Tokyo."

# Load the GPT-2 model and generate answer
generator = pipeline('question-answering', model='distilgpt2')
answer = generator(question=question, context=context)

# Print the answer
print(answer['answer'])

def speak_time():
    # Get the current time in hours, minutes, and AM/PM format
    current_time = datetime.datetime.now().strftime("%I:%M %p")
    # Get the current date and day of the week
    current_date = datetime.date.today().strftime("%B %d, %Y")
    current_day = datetime.date.today().strftime("%A")
    # Print the current time, date, and day of the week to the console
    print("sahi ni saa:", current_time)
    print("Tarehe ni:", current_date)
    print("Siku ni:", current_day)

    # Speak the current time, date, and day of the week using the speech function
    speech(f"sahi ni saa {current_time}, leo ni {current_day}, na tarehe ni {current_date}.")


def get_location():
    geolocator = Nominatim(user_agent="geoapiExercises")
    ip_request = requests.get('https://get.geojs.io/v1/ip.json')
    my_ip = ip_request.json()['ip']
    geo_request_url = 'https://get.geojs.io/v1/ip/geo/' + my_ip + '.json'
    geo_request = requests.get(geo_request_url)
    geo_data = geo_request.json()
    location = geolocator.reverse(f"{geo_data['latitude']}, {geo_data['longitude']}")
    city = location.raw['address'].get('city', '')
    state = location.raw['address'].get('state', '')
    country = location.raw['address'].get('country', '')
    latitude = geo_data['latitude']
    longitude = geo_data['longitude']
    return city, state, country, latitude, longitude

# Set up API parameters
origin = "New York, NY"
destination = "Los Angeles, CA"
api_key = "YOUR_API_KEY"

# Send request to Google Maps API
url = f"https://maps.googleapis.com/maps/api/directions/json?origin={origin}&destination={destination}&key={api_key}"
response = requests.get(url)

# Parse response JSON to get directions
if response.status_code == 200:
    directions = response.json()["routes"][0]["legs"][0]["steps"]
    for step in directions:
        print(step["html_instructions"])
else:
    print("Error getting directions:", response.status_code)

# Set up navigation command
destination = "New York, NY"
navigation_command = f"googlemaps://?daddr={destination}"

# Open Google Maps app and start navigation
subprocess.call(navigation_command, shell=True)
def open_folder(folder_name):
    path = os.path.join(os.getcwd(), folder_name)
    if os.path.exists(path):
        os.startfile(path)
        return f"nafungua folda {folder_name}..."
    else:
        return f"Folda {folder_name} haipatikanit."

# Define a function to listen for the "listen" keyword and take notes

recoder = sr.Recognizer()
def listen_and_take_notes():
    with sr.Microphone() as source:
        print("Say something!")
        audio = recoder.listen(source)

    try:
        text = recoder.recognize_google(audio)
        if "listen" in text.lower():
            # Create a new Word document if it doesn't exist
            if not os.path.exists("notes.docx"):
                doc = docx.Document()
                doc.save("notes.docx")
            # Open the existing Word document and append the new notes
            doc = docx.Document("notes.docx")
            doc.add_paragraph(text)
            doc.save("notes.docx")
            print("Notes saved successfully!")
    except sr.UnknownValueError:
        print("Speech recognition could not understand audio")
    except sr.RequestError as e:
        print("Could not request results from Speech Recognition service; {0}".format(e))

# Call the function to start listening for the "listen" keyword
listen_and_take_notes()


# Define a function to set a reminder
def set_reminder():
    # Get the current date and time
    now = datetime.datetime.now()

    # Set the reminder time to 10 seconds from now
    reminder_time = now + datetime.timedelta(seconds=10)

    print(f"Reminder set for {reminder_time.strftime('%H:%M:%S')}")

    # Wait until the reminder time
    time.sleep((reminder_time - now).total_seconds())

    print("Reminder: Don't forget to take a break!")


# Define a function to set a timer
def set_timer():
    # Get the timer duration from the user
    duration = int(input("Enter the timer duration in seconds: "))

    # Wait for the specified duration
    time.sleep(duration)

    print("Timer finished!")


# Define a function to set an alarm
def set_alarm():
    # Get the alarm time from the user
    alarm_time = input("Enter the alarm time in HH:MM format: ")

    # Convert the alarm time to a datetime object
    alarm_time = datetime.datetime.strptime(alarm_time, "%H:%M")

    # Get the current date and time
    now = datetime.datetime.now()

    # Calculate the time until the alarm
    time_until_alarm = datetime.datetime.combine(now.date(), alarm_time.time()) - now

    # Wait until the alarm time
    time.sleep(time_until_alarm.total_seconds())

    print("Alarm!")


# Call the functions to set the reminder, timer, and alarm
set_reminder()
set_timer()
set_alarm()

# Define the credentials for accessing the Google Calendar API
SCOPES = ['https://www.googleapis.com/auth/calendar']
SERVICE_ACCOUNT_FILE = 'path/to/service_account.json'
credentials = service_account.Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=SCOPES)

# Create a connection to the Google Calendar API
service = build('calendar', 'v3', credentials=credentials)

# Set the start and end times for the event
event_start = datetime.datetime(2023, 3, 24, 10, 0, tzinfo=pytz.UTC)
event_end = datetime.datetime(2023, 3, 24, 11, 0, tzinfo=pytz.UTC)

# Define the event details
event = {
    'summary': 'Meeting with John',
    'location': 'Virtual',
    'description': 'Discuss project proposal',
    'start': {
        'dateTime': event_start.isoformat(),
        'timeZone': 'UTC',
    },
    'end': {
        'dateTime': event_end.isoformat(),
        'timeZone': 'UTC',
    },
    'reminders': {
        'useDefault': True,
    },
}

# Insert the event into the calendar
event = service.events().insert(calendarId='primary', body=event).execute()
print(f'Event created: {event.get("htmlLink")}')


# Twilio account information
account_sid = 'your_account_sid'
auth_token = 'your_auth_token'

# Twilio phone number
twilio_number = '+1234567890'

# Your phone number
your_number = '+0987654321'

# Create a Twilio client
client = Client(account_sid, auth_token)

# Make a phone call
call = client.calls.create(
    to=your_number,
    from_=twilio_number,
    url='http://demo.twilio.com/docs/voice.xml'
)

print(call.sid)

# Define Twilio credentials and message details
account_sid = "your_account_sid"
auth_token = "your_auth_token"
sender_phone_number = "+1234567890"
receiver_phone_number = "+0987654321"
message = "Hello, this is a test message from Python!"

# Connect to the Twilio API
client = Client(account_sid, auth_token)

# Send the text message
message = client.messages.create(
    body=message,
    from_=sender_phone_number,
    to=receiver_phone_number
)

print("Text message sent successfully!")


def send_email():
    # Set up the SMTP server
    smtp_server = "smtp.gmail.com"
    smtp_port = 587
    smtp_username = "your_email@gmail.com"
    smtp_password = "your_password"
    smtp_sender = "your_email@gmail.com"

    # Set up the email message
    email_recipient = "recipient_email@example.com"
    email_subject = "Test Email"
    email_body = "Hello, this is a test email from my virtual assistant!"

    # Connect to the SMTP server and send the email
    with smtplib.SMTP(smtp_server, smtp_port) as server:
        server.starttls()
        server.login(smtp_username, smtp_password)
        server.sendmail(smtp_sender, email_recipient, f"Subject: {email_subject}\n\n{email_body}")

    print("Email sent successfully!")

# Set up authentication credentials
scope = "user-modify-playback-state"
sp = spotipy.Spotify(auth_manager=SpotifyOAuth(scope=scope))

# Search for a track
track_name = "Bohemian Rhapsody"
results = sp.search(q=track_name, limit=1, type="track")
track_uri = results["tracks"]["items"][0]["uri"]

# Play the track on the active device
devices = sp.devices()["devices"]
if len(devices) > 0:
    device_id = devices[0]["id"]
    sp.start_playback(device_id=device_id, uris=[track_uri])
    print("Playing: ", track_name)
else:
    print("No active devices found")

# Set up the AWS IoT client
iot_client = boto3.client('iot-data')

# Define the device name and action
device_name = 'my_smart_device'
action = 'turn_on'

# Create a payload with the device name and action
payload = {
    'device': device_name,
    'action': action,
}

# Publish the payload to the AWS IoT topic for device control
iot_client.publish(
    topic='smart_home/device_control',
    payload=json.dumps(payload)
)

# Wait for the device to respond with a confirmation message
response = iot_client.subscribe(
    topic='smart_home/device_response',
    qos=1
)

# Parse the response message to ensure the device executed the action successfully
message = response['payload']
if message.get('status') == 'success':
    print(f'{device_name} was successfully turned {action}!')
else:
    print(f'{device_name} failed to turn {action}.')

def calculate(operation, num1, num2):
    if operation == 'add':
        result = num1 + num2
    elif operation == 'subtract':
        result = num1 - num2
    elif operation == 'multiply':
        result = num1 * num2
    elif operation == 'divide':
        result = num1 / num2
    else:
        print('Invalid operation')
        return None
    return result

def recommend_music(mood):
    if mood == 'happy':
        songs = ['Happy - Pharrell Williams', 'Can’t Stop The Feeling - Justin Timberlake', 'Good Time - Carly Rae Jepsen, Owl City']
    elif mood == 'sad':
        songs = ['All I Want - Kodaline', 'Fix You - Coldplay', 'Say You Love Me - Jessie Ware']
    elif mood == 'chill':
        songs = ['Lost In Japan - Shawn Mendes', 'Sunflower - Post Malone, Swae Lee', 'Better Together - Jack Johnson']
    else:
        print('Invalid mood')
        return None
    return songs
songs = recommend_music('happy')
print(songs)  # Output: ['Happy - Pharrell Williams', 'Can’t Stop The Feeling - Justin Timberlake', 'Good Time - Carly Rae Jepsen, Owl City']

songs = recommend_music('sad')
print(songs)  # Output: ['All I Want - Kodaline', 'Fix You - Coldplay', 'Say You Love Me - Jessie Ware']

songs = recommend_music('angry')
print(songs)  # Output: Invalid mood

# Set up the request parameters
url = 'https://api.venmo.com/v1/payments'
access_token = 'your_access_token_here'
data = {
    'access_token': access_token,
    'user_id': 'recipient_user_id_here',
    'amount': '10.00',
    'note': 'Thanks for lunch!',
}

# Send the request
response = requests.post(url, data=data)

# Check the response status code
if response.status_code == 200:
    print('Payment sent successfully!')
else:
    print('Error sending payment:', response.text)

#main loop

while True:
    text = get_audio()

    if "video" in text.lower():
        search_query = get_audio()
        speech(f"sawa, nitatafta  {search_query} kwa sababu yako")
        pywhatkit.playonyt(search_query)
    elif "mzaha" in text.lower():
        speech(pyjokes.get_joke())
    elif "nani alikutengeneza" in text.lower():
        speech("nilitengenezwa na mtaalamu wa IT  Leslie Kadenge kutoka chuo kikuu cha Zetech University .")
    elif text.lower() == "rafiki":
        speech("Naaam, Nikusaidie vipi")
    elif "saa" in text.lower():
        speak_time()
    elif "fungua calculator" in text.lower():
        os.startfile("calc.exe")
        speech("nafugua Calculator.")
    elif "fungua browser" in text.lower():
        os.startfile("C:\\Program Files\\BraveSoftware\\Brave-Browser\\Application\\brave.exe")
        speech("nafungua Browser.")
    elif "fungua file explorer" in text.lower():
        os.startfile("explorer.exe")
        speech("nafungua File Explorer.")
    elif "fungua notepad" in text.lower():
        os.startfile("C:\\Program Files (x86)\\Notepad++\\notepad++.exe")
        speech("nafungua notepad+.")
    elif "fungua spotify" in text.lower():
        os.startfile("C:\\Users\\admin\\AppData\\Roaming\\Spotify\\Spotify.exe")
        speech("nafungua spotify.")
    elif "fungua telegram" in text.lower():
        os.startfile("C:\\Users\\admin\\AppData\\Roaming\\Telegram Desktop\\Telegram.exe")
        speech("nafungua telegram.")
    elif "discord" in text.lower():
        os.startfile("C:\\Users\\admin\\AppData\\Local\\Discord\\Update.exe --processStart Discord.exe")
        speech("nafungua discord.")
    elif "fungua epic games" in text.lower():
        os.startfile("C:\\Program Files (x86)\\Epic Games\\Launcher\\Portal\\Binaries\\Win32\\EpicGamesLauncher.exe")
        speech("nafungua epic games.")
    elif "fungua binance" in text.lower():
        os.startfile("C:\\Program Files\\Binance\\Binance.exe")
        speech("nafungua binance.")
    elif "nafungua visual studio" in text.lower():
        os.startfile("C:\\Users\\admin\\AppData\\Local\\Programs\\Microsoft VS Code\\Code.exe")
        speech("fungua vscode.")
    elif "fungua pycharm" in text.lower():
        os.startfile("C:\\Program Files\\JetBrains\\PyCharm Community Edition 2022.3.2\\bin\pycharm64.exe")
        speech("nafungua pycharm.")
    elif "fungua iq option" in text.lower():
        os.startfile("C:\\Program Files (x86)\\IQ Option\\IQ Option.exe")
        speech("nafungua iq option.")
    elif "fungua oracle vm" in text.lower():
        os.startfile("C:\\Program Files\\Oracle\\VirtualBox\\VirtualBox.exe")
        speech("nafungua oracle vm.")
    elif "fungua inSSIDer" in text.lower():
        os.startfile("C:\\Users\\admin\\AppData\\Local\\inSSIDer\\inSSIDer.exe")
        speech("nafungua inSSIDer.")
    elif "fungua android studio" in text.lower():
        os.startfile("C:\\Program Files\\Android\\Android Studio\\bin\\studio64.exe")
        speech("nafungua android studio.")
    elif "fungua expert option" in text.lower():
        os.startfile("C:\\Program Files\\ExpertOption\\ExpertOption.exe")
        speech("fungua expert option.")
    elif "fungua whatsapp" in text.lower():
        pywhatkit.sendwhatmsg_instantly(phone_no="+25494383225", message="/app")
        speech("nafungua WhatsApp.")
    elif "fungua camera" in text.lower():
        pywhatkit.sendwhatmsg_instantly(phone_no="+254794383225", message="/camera")
        speech("nafungua Camera.")
    elif "fungua instagram" in text.lower():
        pywhatkit.sendwhatmsg_instantly(phone_no="+254794383225", message="https://www.instagram.com/")
        speech("nafungua Instagram.")
    elif "fungua twitter" in text.lower():
        pywhatkit.sendwhatmsg_instantly(phone_no="+254794383225", message="https://twitter.com/home")
        speech("fungua Twitter.")
    elif "nafungua tiktok" in text.lower():
        pywhatkit.sendwhatmsg_instantly(phone_no="+254794383225", message="https://www.tiktok.com/")
        speech("nafungua TikTok.")
    elif "fungua facebook" in text.lower():
        pywhatkit.sendwhatmsg_instantly(phone_no="+25494383225", message="https://www.facebook.com/")
        speech("fungua Facebook.")
    elif "niko wapi" in text.lower() or "hapa ni wapi" in text.lower():
        city, state, country, latitude, longitude = get_location()
        speech(f"kwa sasa uko  {city}, {state}, {country}. latitudo {latitude} na longitudo {longitude}.")
        pywhatkit.search(f"My location is {city}, {state}, {country} on Google Maps")
    elif 'hali ya hewa' in text.lower():
        if 'in' in text.lower():
            city = text.lower().split('in')[-1].strip()
        else:
            g = geocoder.ip('me')
            city = f'{g.city}, {g.country}'
        owm = pyowm.OWM('44b31c6431095913f312e5ab9b5d3237')  # replace with your own API key
        mgr = owm.weather_manager()
        try:
            observation = mgr.weather_at_place(city)
            weather = observation.weather
            temp = weather.temperature('celsius')['temp']
            status = weather.detailed_status
            speech(f'kipimo cha joto katika {city} ni {temp} digrii za Celsius na {status}.')
        except pyowm.exceptions.api_response_error.NotFoundError:
            speech(f"Samahani, zijaweza kupata hali ya hewa ya  {city}.")
    elif 'fungua folda' in text:
        # Extract the folder name from the user input
        folder_name = text.replace('open folder ', '')

        # Define the path to the folder based on the name given by the user
        folder_path = os.path.join(os.path.expanduser('~'), folder_name)

        try:
            # Open the folder in the file explorer
            os.startfile(folder_path)
            speech(f"nafungua {folder_name} folda")
        except FileNotFoundError:
            speech(f"samahani, sijaweza kupata folder {folder_name}")

    else:
        query = text.lower()
        results = googlesearch.search(query, num_results=1)
        for result in results:
            response = requests.get(result)
            soup = BeautifulSoup(response.content, 'html.parser')
            text = soup.get_text()
            speech(f"Hiki ndicho nilichopata: {text}")


# Check if the user said "thanks"

















